from docx import Document

def save_resume_to_docx(resume_text, filename="formatted_resume.docx"):
    doc = Document()

    # Split text into lines
    lines = resume_text.split("\n")

    # Add a title
    doc.add_heading("Resume", level=1)

    for line in lines:
        line = line.strip()
        if not line:
            continue  # Skip empty lines

        # Identify section headings and apply formatting
        if "**" in line:  
            clean_line = line.replace("**", "").strip()
            if clean_line.isupper():  # Major sections like "SUMMARY", "EXPERIENCE"
                doc.add_heading(clean_line, level=2)
            else:  # Subheadings like job titles
                doc.add_heading(clean_line, level=3)

        elif line.startswith("* "):  # Bullet points for experience, skills
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)  # Normal paragraph

    # Save the formatted resume
    doc.save(filename)
    print(f"Resume saved as {filename}")